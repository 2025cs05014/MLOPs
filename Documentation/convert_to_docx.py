"""
Convert the Markdown report to a formatted .docx file.
Usage: python Documentation/convert_to_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).resolve().parent.parent
MD_PATH = Path(__file__).resolve().parent / "Heart_Disease_MLOps_Report.md"
DOCX_PATH = Path(__file__).resolve().parent / "Heart_Disease_MLOps_Report.docx"
SCREENSHOTS_DIR = BASE_DIR / "screenshots"


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_table(doc, header_row, data_rows):
    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(header_row):
        set_cell_text(table.rows[0].cells[i], h.strip(), bold=True, size=9)

    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            if i < num_cols:
                set_cell_text(row.cells[i], val.strip(), size=9)

    return table


def try_add_image(doc, image_name, width_inches=5.5):
    img_path = SCREENSHOTS_DIR / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def parse_and_build():
    doc = Document()

    # -- Default style --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = MD_PATH.read_text().splitlines()

    i = 0
    in_code_block = False
    code_lines = []
    skip_table_separator = False

    # Collect tables
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # --- Code blocks ---
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                # Flush any table buffer
                if in_table and table_buffer:
                    _flush_table(doc, table_buffer)
                    table_buffer = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Horizontal rules ---
        if line.strip() == "---":
            i += 1
            continue

        # --- Tables ---
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Check if it's a separator row (e.g., |---|---|)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_buffer.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table and table_buffer:
                _flush_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        # --- Headings ---
        if line.startswith("# ") and not line.startswith("## "):
            text = _clean_md(line[2:])
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## "):
            text = _clean_md(line[3:])
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("### "):
            text = _clean_md(line[4:])
            doc.add_heading(text, level=2)
            i += 1
            continue

        # --- Bold lines (like **Student ID:**) ---
        if line.strip().startswith("**") and ":**" in line:
            p = doc.add_paragraph()
            _add_formatted_run(p, line.strip())
            i += 1
            continue

        # --- Bullet points ---
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text = _clean_md(bullet_match.group(2))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            _add_formatted_run(p, text)
            i += 1
            continue

        # --- Empty lines ---
        if line.strip() == "":
            i += 1
            continue

        # --- Normal paragraph ---
        text = _clean_md(line.strip())
        if text:
            p = doc.add_paragraph()
            _add_formatted_run(p, text)

        i += 1

    # Flush remaining table
    if in_table and table_buffer:
        _flush_table(doc, table_buffer)

    # --- Add screenshots as figures ---
    doc.add_page_break()
    doc.add_heading("Appendix: Screenshots & Visualizations", level=1)

    screenshot_items = [
        ("class_balance.png", "Target Class Distribution"),
        ("feature_histograms.png", "Feature Distributions by Target"),
        ("categorical_features.png", "Categorical Features vs Target"),
        ("correlation_heatmap.png", "Feature Correlation Heatmap"),
        ("boxplots.png", "Box Plots — Outlier Detection"),
        ("cm_logistic_regression.png", "Confusion Matrix — Logistic Regression"),
        ("cm_random_forest.png", "Confusion Matrix — Random Forest"),
        ("roc_logistic_regression.png", "ROC Curve — Logistic Regression"),
        ("roc_random_forest.png", "ROC Curve — Random Forest"),
    ]

    for filename, caption in screenshot_items:
        if try_add_image(doc, filename, width_inches=4.5):
            p = doc.add_paragraph(f"Figure: {caption}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].italic = True
            doc.add_paragraph()  # spacing

    doc.save(str(DOCX_PATH))
    print(f"Report saved to {DOCX_PATH}")


def _flush_table(doc, table_buffer):
    if len(table_buffer) < 2:
        return
    header = table_buffer[0]
    rows = table_buffer[1:]
    add_table(doc, header, rows)
    doc.add_paragraph()  # spacing


def _clean_md(text):
    # Remove markdown link syntax [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove backticks (we'll handle inline code as plain text in docx)
    text = text.replace('`', '')
    return text.strip()


def _add_formatted_run(paragraph, text):
    """Parse bold (**text**) and add runs accordingly."""
    paragraph.clear()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


if __name__ == "__main__":
    parse_and_build()
